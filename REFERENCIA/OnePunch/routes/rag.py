"""
RAG Document Management Routes
"""
import os
import uuid
from datetime import datetime
from concurrent.futures import ThreadPoolExecutor
from flask import Blueprint, request, jsonify, current_app, copy_current_request_context
from flask_jwt_extended import jwt_required, get_jwt_identity
from werkzeug.utils import secure_filename
from extensions import db
from models.user import User
from models.document import Document, DocumentType, DocumentStatus

rag_bp = Blueprint('rag', __name__)

# Thread pool for background document processing
# max_workers=3 limits concurrent processing to avoid overloading
document_executor = ThreadPoolExecutor(max_workers=3, thread_name_prefix='rag_processor')


def get_user_company():
    """Helper to get current user's company"""
    try:
        user_id = int(get_jwt_identity())
        user = User.query.get(user_id)
        if not user:
            print(f"DEBUG: RAG - User {user_id} not found")
            return None
        if not user.company:
            print(f"DEBUG: RAG - User {user.id} has no company")
            return None
        return user.company
    except Exception as e:
        print(f"DEBUG: RAG - Error getting company: {e}")
        return None


def allowed_file(filename):
    """Check if file extension is allowed"""
    return '.' in filename and \
           filename.rsplit('.', 1)[1].lower() in current_app.config.get('ALLOWED_EXTENSIONS', {'pdf', 'docx', 'txt'})


@rag_bp.route('/documents', methods=['GET'])
@jwt_required()
def list_documents():
    """List all documents for the company"""
    company = get_user_company()
    if not company:
        return jsonify({'error': 'Company not found'}), 404
    
    documents = Document.query.filter_by(company_id=company.id).order_by(Document.uploaded_at.desc()).all()
    return jsonify({'documents': [d.to_dict() for d in documents]})


def extract_text_content(file_path: str, file_type: str) -> str:
    """
    Extract text content from file.
    Supports: TXT, CSV, PDF, DOCX
    Uses lazy imports to minimize memory usage.
    """
    try:
        if file_type == 'txt':
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                return f.read()
                
        elif file_type == 'csv':
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                return f.read()
                
        elif file_type == 'pdf':
            # Lazy import pdfminer only when needed
            from pdfminer.high_level import extract_text
            return extract_text(file_path)
            
        elif file_type == 'docx':
            # Lazy import python-docx only when needed
            from docx import Document as DocxDocument
            doc = DocxDocument(file_path)
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            return '\n'.join(paragraphs)
            
        else:
            return None
            
    except Exception as e:
        print(f"Error extracting text from {file_type}: {e}")
        return None


def process_document_sync(document_id: int, company_id: int, app):
    """
    Process document and index to Pinecone - Background thread safe.
    Creates its own app context for database operations.
    
    Args:
        document_id: ID of the document to process
        company_id: ID of the company (for RAG namespace)
        app: Flask application instance for context
    """
    from services.rag_service import RAGService
    from models.company import Company
    
    with app.app_context():
        try:
            # Fetch fresh instances within this context
            document = Document.query.get(document_id)
            company = Company.query.get(company_id)
            
            if not document or not company:
                print(f"Document {document_id} or Company {company_id} not found")
                return False
            
            # Update status to processing
            document.status = DocumentStatus.PROCESSING
            db.session.commit()
            
            # Extract text content
            content = extract_text_content(document.file_path, document.file_type.value)
            
            if not content:
                document.status = DocumentStatus.ERROR
                document.error_message = f"Cannot extract text from {document.file_type.value} files. Supported: TXT, CSV, PDF, DOCX."
                db.session.commit()
                return False
            
            if len(content.strip()) < 10:
                document.status = DocumentStatus.ERROR
                document.error_message = "Document content is too short"
                db.session.commit()
                return False
            
            # Initialize RAG service with company context
            rag_service = RAGService(company)
            
            # Index document to Pinecone
            metadata = {
                'title': document.title,
                'filename': document.original_filename,
                'file_type': document.file_type.value
            }
            
            vector_ids = rag_service.index_document(
                document_id=document.id,
                content=content,
                metadata=metadata
            )
            
            # Update document record
            document.pinecone_ids = vector_ids
            document.chunk_count = len(vector_ids)
            document.status = DocumentStatus.INDEXED
            document.processed_at = datetime.utcnow()
            document.error_message = None
            db.session.commit()
            
            print(f"✓ Document {document_id} indexed successfully ({len(vector_ids)} chunks)")
            return True
            
        except Exception as e:
            try:
                document = Document.query.get(document_id)
                if document:
                    document.status = DocumentStatus.ERROR
                    document.error_message = str(e)[:500]
                    db.session.commit()
            except:
                pass
            print(f"✗ Error processing document {document_id}: {e}")
            return False


@rag_bp.route('/upload', methods=['POST'])
@jwt_required()
def upload_document():
    """Upload a document for RAG processing (async)"""
    company = get_user_company()
    if not company:
        return jsonify({'error': 'Company not found'}), 404
    
    if 'file' not in request.files:
        return jsonify({'error': 'No file provided'}), 400
    
    file = request.files['file']
    
    if file.filename == '':
        return jsonify({'error': 'No file selected'}), 400
    
    if not allowed_file(file.filename):
        return jsonify({'error': 'File type not allowed'}), 400
    
    try:
        # Generate unique filename
        original_filename = secure_filename(file.filename)
        extension = original_filename.rsplit('.', 1)[1].lower()
        unique_filename = f"{uuid.uuid4()}.{extension}"
        
        # Create company upload directory
        upload_dir = os.path.join(current_app.config['UPLOAD_FOLDER'], str(company.id))
        os.makedirs(upload_dir, exist_ok=True)
        
        file_path = os.path.join(upload_dir, unique_filename)
        file.save(file_path)
        
        # Get file size
        file_size = os.path.getsize(file_path)
        
        # Create document record
        document = Document(
            company_id=company.id,
            filename=unique_filename,
            original_filename=original_filename,
            file_type=DocumentType(extension),
            file_size=file_size,
            file_path=file_path,
            title=request.form.get('title', original_filename),
            description=request.form.get('description'),
            status=DocumentStatus.PENDING
        )
        db.session.add(document)
        db.session.commit()
        
        # Get IDs and app for background thread
        doc_id = document.id
        company_id = company.id
        app = current_app._get_current_object()
        
        # Submit to background thread pool (non-blocking)
        document_executor.submit(process_document_sync, doc_id, company_id, app)
        
        return jsonify({
            'message': 'Document uploaded, processing in background',
            'document': document.to_dict()
        }), 201
        
    except Exception as e:
        db.session.rollback()
        return jsonify({'error': str(e)}), 500


@rag_bp.route('/documents/<int:document_id>', methods=['GET'])
@jwt_required()
def get_document(document_id):
    """Get document details"""
    company = get_user_company()
    if not company:
        return jsonify({'error': 'Company not found'}), 404
    
    document = Document.query.filter_by(id=document_id, company_id=company.id).first()
    if not document:
        return jsonify({'error': 'Document not found'}), 404
    
    return jsonify({'document': document.to_dict()})


@rag_bp.route('/documents/<int:document_id>', methods=['PUT'])
@jwt_required()
def update_document(document_id):
    """Update document metadata"""
    company = get_user_company()
    if not company:
        return jsonify({'error': 'Company not found'}), 404
    
    document = Document.query.filter_by(id=document_id, company_id=company.id).first()
    if not document:
        return jsonify({'error': 'Document not found'}), 404
    
    data = request.get_json()
    
    if 'title' in data:
        document.title = data['title']
    if 'description' in data:
        document.description = data['description']
    
    db.session.commit()
    
    return jsonify({
        'message': 'Document updated',
        'document': document.to_dict()
    })


@rag_bp.route('/documents/<int:document_id>', methods=['DELETE'])
@jwt_required()
def delete_document(document_id):
    """Delete a document"""
    company = get_user_company()
    if not company:
        return jsonify({'error': 'Company not found'}), 404
    
    document = Document.query.filter_by(id=document_id, company_id=company.id).first()
    if not document:
        return jsonify({'error': 'Document not found'}), 404
    
    try:
        # Delete file
        if document.file_path and os.path.exists(document.file_path):
            os.remove(document.file_path)
        
        # Delete vectors from Pinecone
        try:
            from services.rag_service import RAGService
            rag_service = RAGService(company)
            rag_service.delete_document(document.id)
        except Exception as e:
            print(f"Warning: Could not delete vectors from Pinecone: {e}")
        
        db.session.delete(document)
        db.session.commit()
        
        return jsonify({'message': 'Document deleted'})
        
    except Exception as e:
        db.session.rollback()
        return jsonify({'error': str(e)}), 500


@rag_bp.route('/documents/<int:document_id>/reprocess', methods=['POST'])
@jwt_required()
def reprocess_document(document_id):
    """Reprocess a document for RAG"""
    company = get_user_company()
    if not company:
        return jsonify({'error': 'Company not found'}), 404
    
    document = Document.query.filter_by(id=document_id, company_id=company.id).first()
    if not document:
        return jsonify({'error': 'Document not found'}), 404
    
    # Mark as pending for reprocessing
    document.status = DocumentStatus.PENDING
    document.error_message = None
    db.session.commit()
    
    # Get IDs and app for background thread
    doc_id = document.id
    company_id = company.id
    app = current_app._get_current_object()
    
    # Submit to background thread pool (non-blocking)
    document_executor.submit(process_document_sync, doc_id, company_id, app)
    
    return jsonify({
        'message': 'Document queued for reprocessing',
        'document': document.to_dict()
    })


@rag_bp.route('/query', methods=['POST'])
@jwt_required()
def query_documents():
    """Query RAG memory"""
    company = get_user_company()
    if not company:
        return jsonify({'error': 'Company not found'}), 404
    
    data = request.get_json()
    query = data.get('query', '')
    top_k = data.get('top_k', 5)
    
    if not query:
        return jsonify({'error': 'Query is required'}), 400
    
    try:
        from services.rag_service import RAGService
        rag_service = RAGService(company)
        results = rag_service.query(query, top_k=top_k)
        
        return jsonify({
            'query': query,
            'results': results,
            'count': len(results)
        })
    except Exception as e:
        return jsonify({
            'query': query,
            'results': [],
            'error': str(e)
        })


@rag_bp.route('/stats', methods=['GET'])
@jwt_required()
def get_rag_stats():
    """Get RAG statistics"""
    company = get_user_company()
    if not company:
        return jsonify({'error': 'Company not found'}), 404
    
    total_docs = Document.query.filter_by(company_id=company.id).count()
    indexed_docs = Document.query.filter_by(company_id=company.id, status=DocumentStatus.INDEXED).count()
    pending_docs = Document.query.filter_by(company_id=company.id, status=DocumentStatus.PENDING).count()
    error_docs = Document.query.filter_by(company_id=company.id, status=DocumentStatus.ERROR).count()
    
    total_chunks = db.session.query(db.func.sum(Document.chunk_count)).filter_by(company_id=company.id).scalar() or 0
    
    return jsonify({
        'total_documents': total_docs,
        'indexed_documents': indexed_docs,
        'pending_documents': pending_docs,
        'error_documents': error_docs,
        'total_chunks': total_chunks
    })
